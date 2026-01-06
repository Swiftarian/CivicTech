# -*- coding: utf-8 -*-
"""
差異報告產生模組 - 生成 Word/PDF 格式的比對報告
"""

import os
from datetime import datetime
from typing import List, Dict, Optional
from pathlib import Path

# 延遲導入
def get_docx():
    try:
        from docx import Document
        from docx.shared import Pt, Cm, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        return Document, Pt, Cm, RGBColor, WD_ALIGN_PARAGRAPH, qn
    except ImportError:
        raise ImportError("請安裝 python-docx: pip install python-docx")


def create_comparison_report(comparison_results: List[Dict],
                            output_path: str,
                            title: str = "文件資料比對報告",
                            include_summary: bool = True) -> str:
    """
    生成 Word 格式的比對報告
    
    Args:
        comparison_results: 比對結果列表
        output_path: 輸出檔案路徑
        title: 報告標題
        include_summary: 是否包含摘要統計
    
    Returns:
        輸出檔案路徑
    """
    Document, Pt, Cm, RGBColor, WD_ALIGN_PARAGRAPH, qn = get_docx()
    
    doc = Document()
    
    # 設定頁面邊界
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # 標題
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 報告生成時間
    timestamp = datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")
    time_para = doc.add_paragraph(f"報告生成時間：{timestamp}")
    time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # 空行
    
    # 摘要統計
    if include_summary:
        doc.add_heading("摘要統計", level=1)
        
        total = len(comparison_results)
        passed = sum(1 for r in comparison_results if r.get("overall_match", False))
        failed = total - passed
        avg_similarity = sum(r.get("overall_similarity", 0) for r in comparison_results) / total if total > 0 else 0
        
        summary_table = doc.add_table(rows=4, cols=2)
        summary_table.style = 'Table Grid'
        
        summary_data = [
            ("比對文件總數", str(total)),
            ("通過數量", f"{passed} ✅"),
            ("未通過數量", f"{failed} ❌"),
            ("平均相似度", f"{avg_similarity:.1%}")
        ]
        
        for i, (label, value) in enumerate(summary_data):
            row = summary_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            
            # 設定字型
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "標楷體"
                        run.font.size = Pt(12)
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
        
        doc.add_paragraph()
    
    # 詳細比對結果
    doc.add_heading("詳細比對結果", level=1)
    
    for idx, result in enumerate(comparison_results, 1):
        doc_id = result.get("document_id", f"文件 {idx}")
        overall_match = result.get("overall_match", False)
        overall_sim = result.get("overall_similarity", 0)
        summary = result.get("summary", "")
        
        # 文件標題
        status_icon = "✅" if overall_match else "❌"
        doc.add_heading(f"{idx}. {doc_id} {status_icon}", level=2)
        
        # 整體結果
        para = doc.add_paragraph()
        para.add_run(f"整體相似度：{overall_sim:.1%}").bold = True
        doc.add_paragraph(summary)
        
        # 欄位比對表格
        field_comparisons = result.get("field_comparisons", [])
        if field_comparisons:
            table = doc.add_table(rows=len(field_comparisons) + 1, cols=5)
            table.style = 'Table Grid'
            
            # 表頭
            headers = ["欄位名稱", "OCR 值", "參考值", "相似度", "結果"]
            header_row = table.rows[0]
            for i, header in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.name = "標楷體"
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
            
            # 資料列
            for row_idx, fc in enumerate(field_comparisons, 1):
                row = table.rows[row_idx]
                
                row.cells[0].text = str(fc.get("field_name", ""))
                row.cells[1].text = str(fc.get("ocr_value", ""))[:50]  # 截斷過長內容
                row.cells[2].text = str(fc.get("reference_value", ""))[:50]
                row.cells[3].text = f"{fc.get('similarity', 0):.1%}"
                
                match_type = fc.get("match_type", "")
                if match_type == "exact":
                    row.cells[4].text = "✅ 完全相符"
                elif match_type == "similar":
                    row.cells[4].text = "⚠️ 相似"
                elif match_type == "mismatch":
                    row.cells[4].text = "❌ 不符"
                elif match_type == "missing":
                    row.cells[4].text = "⚠️ 缺失"
                else:
                    row.cells[4].text = match_type
                
                # 設定字型
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = "標楷體"
                            run.font.size = Pt(10)
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
        
        doc.add_paragraph()  # 間隔
    
    # 頁尾
    doc.add_paragraph()
    footer = doc.add_paragraph("─" * 40)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer = doc.add_paragraph("本報告由 OCR 文件比對系統自動產生")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 儲存
    doc.save(output_path)
    return output_path


def create_simple_report(comparison_results: List[Dict],
                        output_path: str) -> str:
    """
    生成簡易文字格式報告
    
    Args:
        comparison_results: 比對結果列表
        output_path: 輸出檔案路徑
    
    Returns:
        輸出檔案路徑
    """
    lines = []
    lines.append("=" * 60)
    lines.append("文件資料比對報告")
    lines.append(f"生成時間：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append("=" * 60)
    lines.append("")
    
    # 統計
    total = len(comparison_results)
    passed = sum(1 for r in comparison_results if r.get("overall_match", False))
    lines.append(f"總計：{total} 份文件")
    lines.append(f"通過：{passed} 份")
    lines.append(f"未通過：{total - passed} 份")
    lines.append("")
    lines.append("-" * 60)
    
    # 詳細結果
    for idx, result in enumerate(comparison_results, 1):
        doc_id = result.get("document_id", f"文件 {idx}")
        overall_match = "✅ 通過" if result.get("overall_match", False) else "❌ 未通過"
        overall_sim = result.get("overall_similarity", 0)
        
        lines.append(f"\n[{idx}] {doc_id}")
        lines.append(f"    結果：{overall_match}")
        lines.append(f"    相似度：{overall_sim:.1%}")
        
        for fc in result.get("field_comparisons", []):
            field_name = fc.get("field_name", "")
            match_type = fc.get("match_type", "")
            similarity = fc.get("similarity", 0)
            
            if match_type != "exact":
                lines.append(f"    - {field_name}: {match_type} ({similarity:.1%})")
                lines.append(f"      OCR: {fc.get('ocr_value', '')}")
                lines.append(f"      參考: {fc.get('reference_value', '')}")
    
    lines.append("")
    lines.append("=" * 60)
    lines.append("報告結束")
    
    # 寫入檔案
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(lines))
    
    return output_path


# 測試用主程式
if __name__ == "__main__":
    # 測試資料
    test_results = [
        {
            "document_id": "DOC-001",
            "overall_match": True,
            "overall_similarity": 0.98,
            "summary": "✅ 所有欄位比對通過",
            "field_comparisons": [
                {"field_name": "場所名稱", "ocr_value": "台東縣立體育場", "reference_value": "台東縣立體育場", "match_type": "exact", "similarity": 1.0},
                {"field_name": "地址", "ocr_value": "台東市中華路684號", "reference_value": "台東市中華路684號", "match_type": "exact", "similarity": 1.0},
            ]
        },
        {
            "document_id": "DOC-002",
            "overall_match": False,
            "overall_similarity": 0.75,
            "summary": "❌ 不符欄位: 負責人",
            "field_comparisons": [
                {"field_name": "場所名稱", "ocr_value": "某某商店", "reference_value": "某某商店", "match_type": "exact", "similarity": 1.0},
                {"field_name": "負責人", "ocr_value": "王小名", "reference_value": "王小明", "match_type": "similar", "similarity": 0.8},
            ]
        }
    ]
    
    # 生成報告
    output_txt = "test_report.txt"
    create_simple_report(test_results, output_txt)
    print(f"文字報告已生成：{output_txt}")
    
    try:
        output_docx = "test_report.docx"
        create_comparison_report(test_results, output_docx)
        print(f"Word 報告已生成：{output_docx}")
    except ImportError as e:
        print(f"無法生成 Word 報告：{e}")
